"""切片报告 Word 导出服务

将 SocialSlice 中的 AI 报告（Markdown）转换为格式化的 Word 文档。
"""

from __future__ import annotations

import re
from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from .models import SocialSlice


def generate_slice_report_docx(slice_record: SocialSlice) -> BytesIO:
    """从切片记录生成 Word 文档并返回 BytesIO 流。"""
    result_data = slice_record.result_data or {}
    reports = result_data.get("reports") or {}
    overview_data = (
        result_data.get("layers", {}).get("landscape", {}).get("overview") or {}
    )
    monitor_name = (
        slice_record.monitor.name if slice_record.monitor else "未命名项目"
    )
    slice_name = slice_record.name or f"切片 #{slice_record.id}"
    generated_date = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    doc = Document()
    _setup_styles(doc)

    # 封面
    _add_cover_page(doc, monitor_name, slice_name, generated_date)

    # 概览摘要
    _add_overview_section(doc, overview_data)

    # 3 份报告
    report_defs = [
        ("landscape_report", "行业格局报告"),
        ("topic_report", "话题洞察报告"),
        ("focus_report", "战略诊断报告"),
    ]
    for key, title in report_defs:
        content = _extract_report_content(reports, key)
        if content:
            _add_report_chapter(doc, title, content)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# --------------- 内部辅助 ---------------


def _setup_styles(doc: Document) -> None:
    """配置文档默认字体和段落样式。"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(10.5)
    # 设置中文字体
    rpr = style.element.get_or_add_rPr()
    rpr_fonts = rpr.find(qn("w:rFonts"))
    if rpr_fonts is None:
        rpr_fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rpr_fonts)
    rpr_fonts.set(qn("w:eastAsia"), "微软雅黑")

    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.3


def _add_cover_page(
    doc: Document, monitor_name: str, slice_name: str, date_str: str
) -> None:
    """生成封面页。"""
    for _ in range(6):
        doc.add_paragraph("")

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(monitor_name)
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p_sub.add_run(slice_name)
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p_date.add_run(f"生成日期：{date_str}")
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()


def _add_overview_section(doc: Document, overview: dict) -> None:
    """添加概览摘要表格。"""
    doc.add_heading("概览摘要", level=1)

    metrics = [
        ("原文总量", overview.get("total_volume", "-")),
        ("去重原文数", overview.get("unique_posts", "-")),
        ("总热度", overview.get("total_heat", "-")),
        ("全局 NSR", overview.get("global_nsr", "-")),
        ("有机 NSR", overview.get("organic_nsr", "-")),
        ("推广 NSR", overview.get("promo_nsr", "-")),
    ]

    platform_vol = overview.get("platform_volume") or {}
    for platform, count in platform_vol.items():
        metrics.append((f"平台 - {platform}", count))

    table = doc.add_table(rows=len(metrics), cols=2, style="Light Shading Accent 1")
    for idx, (label, value) in enumerate(metrics):
        table.rows[idx].cells[0].text = str(label)
        table.rows[idx].cells[1].text = str(value)

    doc.add_paragraph("")
    doc.add_page_break()


def _extract_report_content(reports: dict, key: str) -> str:
    """从 reports 字典提取指定报告的 Markdown 内容。"""
    report = reports.get(key)
    if not report:
        return ""
    if isinstance(report, dict):
        return report.get("content", "")
    if isinstance(report, str):
        return report
    return ""


def _add_report_chapter(doc: Document, title: str, markdown: str) -> None:
    """将一份 Markdown 报告转换为 Word 章节。"""
    doc.add_heading(title, level=1)
    _render_markdown(doc, markdown)
    doc.add_page_break()


# --------------- Markdown → Word 行级解析 ---------------

# 匹配 Markdown 标题行
_RE_HEADING = re.compile(r"^(#{1,4})\s+(.*)")
# 匹配无序列表项
_RE_LIST_ITEM = re.compile(r"^(\s*)[*\-+]\s+(.*)")
# 匹配有序列表项
_RE_OL_ITEM = re.compile(r"^(\s*)\d+[.)]\s+(.*)")
# 匹配加粗文本
_RE_BOLD = re.compile(r"\*\*(.+?)\*\*")
# 匹配行内代码
_RE_CODE = re.compile(r"`([^`]+)`")


def _render_markdown(doc: Document, markdown: str) -> None:
    """逐行解析 Markdown 并添加到 Word 文档。"""
    in_code_block = False

    for raw_line in markdown.split("\n"):
        line = raw_line.rstrip()

        # 代码块切换
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            p = doc.add_paragraph(line, style="Normal")
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            continue

        # 空行
        if not line.strip():
            doc.add_paragraph("")
            continue

        # 标题
        m = _RE_HEADING.match(line)
        if m:
            level = min(len(m.group(1)) + 1, 4)  # # → Heading 2 (章节标题已是 H1)
            heading_text = m.group(2).strip()
            doc.add_heading(heading_text, level=level)
            continue

        # 无序列表
        m = _RE_LIST_ITEM.match(line)
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatted_text(p, text)
            continue

        # 有序列表
        m = _RE_OL_ITEM.match(line)
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatted_text(p, text)
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_inline_formatted_text(p, line)


def _add_inline_formatted_text(paragraph, text: str) -> None:
    """解析行内 Markdown 格式（加粗、行内代码）并添加 run。"""
    pos = 0
    # 合并 bold 和 code 的匹配，按位置排序处理
    tokens: list[tuple[int, int, str, str]] = []
    for m in _RE_BOLD.finditer(text):
        tokens.append((m.start(), m.end(), "bold", m.group(1)))
    for m in _RE_CODE.finditer(text):
        tokens.append((m.start(), m.end(), "code", m.group(1)))
    tokens.sort(key=lambda t: t[0])

    for start, end, kind, content in tokens:
        if start > pos:
            paragraph.add_run(text[pos:start])
        run = paragraph.add_run(content)
        if kind == "bold":
            run.bold = True
        elif kind == "code":
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        pos = end

    if pos < len(text):
        paragraph.add_run(text[pos:])
