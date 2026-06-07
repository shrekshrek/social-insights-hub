"""策略报告 Word 导出

按 output_type 渲染对应路径的结果：
- campaign_strategy: Insight → Brand Role → Big Idea（仅 selected 分支）
- market_report:     Agenda Map → Landscape → Strategic Brief
- full_strategy:     Agenda Map → Landscape → Insight → Brand Role → Big Idea → Strategic Brief

复用 analysis/export_docx 的样式设置。
"""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.social_media.analysis.export_docx import _setup_styles

from .models import Strategy


def generate_strategy_docx(strategy: Strategy) -> BytesIO:
    """从策略记录生成 Word 文档，按 output_type 渲染对应路径。

    campaign_strategy 多分支仅导出 selected branch（无则按完成度回退）。
    """
    doc = Document()
    _setup_styles(doc)
    _add_cover(doc, strategy)

    output_type = strategy.output_type or "campaign_strategy"
    has_market = output_type in ("market_report", "full_strategy")
    has_campaign = output_type in ("campaign_strategy", "full_strategy")

    # 媒体侧（market_report / full_strategy）：议程图 + 竞争格局
    if has_market:
        _add_agenda_map_section(doc, strategy.agenda_map_result)
        _add_landscape_section(doc, strategy.landscape_result)

    # 消费者侧（campaign_strategy / full_strategy）：洞察 + 品牌角色 + 创意
    if has_campaign:
        _add_insight_section(doc, strategy.insight_result)
        selected = _pick_export_branch(strategy.brand_strategy_branches)
        branch_label = _format_branch_label(selected)
        _add_brand_role_section(
            doc, (selected or {}).get("brand_role"), branch_label=branch_label
        )
        _add_big_idea_section(
            doc, (selected or {}).get("big_idea"), branch_label=branch_label
        )

    # 战略简报（market_report / full_strategy 终层，综合上游）
    if has_market:
        _add_strategic_brief_section(doc, strategy.strategic_brief_result)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _pick_export_branch(branches: list[dict] | None) -> dict | None:
    """选取要导出的分支：selected=true 优先；否则按完成度回退"""
    if not branches:
        return None
    selected = next(
        (b for b in branches if isinstance(b, dict) and b.get("selected")),
        None,
    )
    if selected:
        return selected
    with_big_idea = next(
        (b for b in branches if isinstance(b, dict) and b.get("big_idea")),
        None,
    )
    if with_big_idea:
        return with_big_idea
    with_brand_role = next(
        (b for b in branches if isinstance(b, dict) and b.get("brand_role")),
        None,
    )
    if with_brand_role:
        return with_brand_role
    return branches[0] if isinstance(branches[0], dict) else None


def _format_branch_label(branch: dict | None) -> str:
    """生成分支标题尾缀，如 "（基于 Tension 1：xxxx）"。"""
    if not branch:
        return ""
    tension_id = branch.get("tension_id")
    summary = (branch.get("tension_summary") or "").strip()
    if tension_id is None and not summary:
        return ""
    if summary:
        return f"（基于 Tension {int(tension_id) + 1 if isinstance(tension_id, int) else '?'}：{summary}）"
    return f"（基于 Tension {int(tension_id) + 1}）"


def _add_cover(doc: Document, strategy: Strategy) -> None:
    """封面页"""
    doc.add_paragraph("")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(strategy.name)
    run.font.size = Pt(24)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("策略报告")
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    p.add_run(f"生成日期: {date_str}")

    doc.add_page_break()


def _add_insight_section(doc: Document, insight: dict | None) -> None:
    """第 1 层 Insight (洞察): Social Tension + Brand Opportunity"""
    doc.add_heading("第 1 层 Insight: 洞察", level=1)

    if not insight:
        doc.add_paragraph("（未完成）")
        return

    # Social Tensions
    tensions = insight.get("social_tensions", [])
    if tensions:
        doc.add_heading("Social Tension", level=2)
        for i, t in enumerate(tensions):
            statement = t.get("statement", "")
            confidence = t.get("confidence", "")
            doc.add_heading(f"Tension {i + 1}: {statement}", level=3)
            if confidence:
                doc.add_paragraph(f"置信度: {confidence}")
            _add_evidence_list(doc, t.get("evidence", []))

    # Brand Opportunities
    opportunities = insight.get("brand_opportunities", [])
    if opportunities:
        doc.add_heading("Brand Opportunity", level=2)
        for i, o in enumerate(opportunities):
            statement = o.get("statement", "")
            doc.add_heading(f"Opportunity {i + 1}: {statement}", level=3)
            related = o.get("related_tensions", [])
            if related:
                doc.add_paragraph(f"关联 Tension: {related}")
            _add_evidence_list(doc, o.get("evidence", []))


def _add_brand_role_section(
    doc: Document,
    brand_role: dict | None,
    *,
    branch_label: str = "",
) -> None:
    """第 2 层 Brand Role (品牌角色): Brand Social Role + Social Strategy"""
    doc.add_heading(f"第 2 层 Brand Role: 品牌角色{branch_label}", level=1)

    if not brand_role:
        doc.add_paragraph("（未完成）")
        return

    # Brand Social Role
    role = brand_role.get("brand_social_role", {})
    if role:
        doc.add_heading("Brand Social Role", level=2)
        if role.get("statement"):
            p = doc.add_paragraph()
            run = p.add_run(role["statement"])
            run.bold = True
            run.font.size = Pt(12)
        if role.get("elaboration"):
            doc.add_paragraph(role["elaboration"])
        _add_evidence_list(doc, role.get("evidence", []))

    # Social Strategy
    social_strategy = brand_role.get("social_strategy", {})
    if social_strategy:
        doc.add_heading("Social Strategy", level=2)
        if social_strategy.get("statement"):
            p = doc.add_paragraph()
            run = p.add_run(social_strategy["statement"])
            run.bold = True
            run.font.size = Pt(12)
        if social_strategy.get("core_message"):
            doc.add_paragraph(f"核心信息: {social_strategy['core_message']}")
        if social_strategy.get("rhythm"):
            doc.add_paragraph(f"传播节奏: {social_strategy['rhythm']}")
        _add_evidence_list(doc, social_strategy.get("evidence", []))


def _add_big_idea_section(
    doc: Document,
    big_idea: dict | None,
    *,
    branch_label: str = "",
) -> None:
    """第 3 层 Big Idea (创意): Big Idea + Content Strategy"""
    doc.add_heading(f"第 3 层 Big Idea: 创意{branch_label}", level=1)

    if not big_idea:
        doc.add_paragraph("（未完成）")
        return

    # Big Idea
    idea = big_idea.get("big_idea", {})
    if idea:
        doc.add_heading("Big Idea", level=2)
        if idea.get("statement"):
            p = doc.add_paragraph()
            run = p.add_run(idea["statement"])
            run.bold = True
            run.font.size = Pt(14)
        if idea.get("elaboration"):
            doc.add_paragraph(idea["elaboration"])
        if idea.get("tension_echo"):
            doc.add_paragraph(f"与核心矛盾的呼应: {idea['tension_echo']}")
        _add_evidence_list(doc, idea.get("evidence", []))

    # Content Strategy
    content = big_idea.get("content_strategy", {})
    if content:
        doc.add_heading("Content Strategy", level=2)
        pillars = content.get("pillars", [])
        for i, pillar in enumerate(pillars):
            doc.add_heading(f"支柱 {i + 1}: {pillar.get('name', '')}", level=3)
            if pillar.get("description"):
                doc.add_paragraph(pillar["description"])
            examples = pillar.get("reference_examples", [])
            if examples:
                doc.add_paragraph("参考案例:")
                for ex in examples:
                    doc.add_paragraph(f"  - {ex}", style="List Bullet")
        _add_evidence_list(doc, content.get("evidence", []))


def _add_voices(doc: Document, voices: list[dict] | None) -> None:
    """代表声音列表（quote / speaker / source / source_tier）"""
    for v in voices or []:
        if not isinstance(v, dict):
            continue
        meta = "、".join(
            x for x in [v.get("speaker"), v.get("source"), v.get("source_tier")] if x
        )
        text = f"“{v.get('quote', '')}”" + (f" —— {meta}" if meta else "")
        doc.add_paragraph(text, style="List Bullet")


def _add_agenda_map_section(doc: Document, agenda: dict | None) -> None:
    """市场报告第 1 层 Agenda Map（媒体议程图）"""
    doc.add_heading("媒体议程图 (Agenda Map)", level=1)
    if not agenda:
        doc.add_paragraph("（未完成）")
        return

    if agenda.get("media_landscape_summary"):
        doc.add_paragraph(agenda["media_landscape_summary"])

    narratives = agenda.get("narrative_map") or []
    if narratives:
        doc.add_heading("叙事地图 (Narrative Map)", level=2)
        for n in narratives:
            doc.add_heading(n.get("theme", ""), level=3)
            if n.get("framing"):
                doc.add_paragraph(f"框架: {n['framing']}")
            meta = []
            if n.get("sentiment") is not None:
                meta.append(f"情感 {n['sentiment']}")
            if n.get("credibility"):
                meta.append(f"可信度 {n['credibility']}")
            if n.get("supporting_sources"):
                meta.append(f"支撑来源 {'、'.join(n['supporting_sources'])}")
            if meta:
                doc.add_paragraph("，".join(meta))
            _add_voices(doc, n.get("representative_voices"))

    battles = agenda.get("agenda_battles") or []
    if battles:
        doc.add_heading("议题战场 (Agenda Battles)", level=2)
        for b in battles:
            doc.add_heading(b.get("contested_topic", ""), level=3)
            for c in b.get("camps") or []:
                tiers = c.get("supporting_tiers")
                tiers_str = (
                    "、".join(tiers) if isinstance(tiers, list) else (tiers or "")
                )
                doc.add_paragraph(
                    f"立场: {c.get('stance', '')}（{tiers_str}）", style="List Bullet"
                )
            if b.get("implication"):
                doc.add_paragraph(f"含义: {b['implication']}")

    gaps = agenda.get("attention_gaps") or []
    if gaps:
        doc.add_heading("注意力盲区 (Attention Gaps)", level=2)
        for g in gaps:
            doc.add_paragraph(
                f"{g.get('topic', '')}：{g.get('why_matters', '')}"
                f"（{g.get('risk_or_opportunity', '')}）",
                style="List Bullet",
            )


def _add_landscape_section(doc: Document, landscape: dict | None) -> None:
    """市场报告第 2 层 Landscape（竞争格局）"""
    doc.add_heading("竞争格局 (Landscape)", level=1)
    if not landscape:
        doc.add_paragraph("（未完成）")
        return

    if landscape.get("competitive_summary"):
        doc.add_paragraph(landscape["competitive_summary"])

    players = landscape.get("players") or []
    if players:
        doc.add_heading("玩家 (Players)", level=2)
        for p in players:
            doc.add_heading(f"{p.get('name', '')}（{p.get('role', '')}）", level=3)
            meta = []
            if p.get("media_sov_pct") is not None:
                meta.append(f"媒体声量 {p['media_sov_pct']}%")
            if p.get("media_sentiment") is not None:
                meta.append(f"媒体情感 {p['media_sentiment']}")
            if meta:
                doc.add_paragraph("，".join(meta))
            if p.get("narrative_position"):
                doc.add_paragraph(f"叙事定位: {p['narrative_position']}")
            for c in p.get("key_claims") or []:
                doc.add_paragraph(c, style="List Bullet")
            _add_voices(doc, p.get("representative_voices"))

    pm = landscape.get("positioning_map") or {}
    positions = pm.get("positions") or []
    if positions:
        doc.add_heading("定位图 (Positioning Map)", level=2)
        xa = pm.get("x_axis") or {}
        ya = pm.get("y_axis") or {}
        doc.add_paragraph(
            f"X 轴: {xa.get('label', '')}（{xa.get('low', '')} ↔ {xa.get('high', '')}）"
        )
        doc.add_paragraph(
            f"Y 轴: {ya.get('label', '')}（{ya.get('low', '')} ↔ {ya.get('high', '')}）"
        )
        for pos in positions:
            doc.add_paragraph(
                f"{pos.get('player', '')}：{pos.get('rationale', '')}",
                style="List Bullet",
            )

    battles = landscape.get("discourse_battles") or []
    if battles:
        doc.add_heading("话语权之争 (Discourse Battles)", level=2)
        for b in battles:
            doc.add_heading(b.get("battle", ""), level=3)
            if b.get("shift_direction"):
                doc.add_paragraph(f"走向: {b['shift_direction']}")
            if b.get("evidence"):
                doc.add_paragraph(f"依据: {b['evidence']}")

    md = landscape.get("market_dynamics") or {}
    if md:
        doc.add_heading("市场动态 (Market Dynamics)", level=2)
        for label, key in (("上升", "momentum_gainers"), ("下降", "momentum_losers")):
            for item in md.get(key) or []:
                doc.add_paragraph(
                    f"{label}: {item.get('player', '')} — {item.get('signal', '')}",
                    style="List Bullet",
                )
        for s in md.get("structural_shifts") or []:
            doc.add_paragraph(
                f"结构变化: {s.get('shift', '')} → {s.get('implication', '')}",
                style="List Bullet",
            )


def _add_strategic_brief_section(doc: Document, brief: dict | None) -> None:
    """市场报告第 3 层 Strategic Brief（战略简报，综合终层）"""
    doc.add_heading("战略简报 (Strategic Brief)", level=1)
    if not brief:
        doc.add_paragraph("（未完成）")
        return

    if brief.get("generation_mode"):
        view = "综合视角" if brief["generation_mode"] == "comprehensive" else "媒体视角"
        doc.add_paragraph(f"视角: {view}")

    if brief.get("executive_summary"):
        doc.add_heading("执行摘要", level=2)
        doc.add_paragraph(brief["executive_summary"])

    priorities = brief.get("strategic_priorities") or []
    if priorities:
        doc.add_heading("战略优先级", level=2)
        for i, p in enumerate(priorities):
            doc.add_heading(f"{i + 1}. {p.get('priority', '')}", level=3)
            if p.get("rationale"):
                doc.add_paragraph(p["rationale"])
            for a in p.get("actions") or []:
                doc.add_paragraph(a, style="List Bullet")
            if p.get("success_metric"):
                doc.add_paragraph(f"成功指标: {p['success_metric']}")

    opps = brief.get("market_opportunities") or []
    if opps:
        doc.add_heading("市场机会", level=2)
        for o in opps:
            doc.add_heading(o.get("opportunity", ""), level=3)
            if o.get("brand_fit"):
                doc.add_paragraph(f"品牌契合: {o['brand_fit']}")
            if o.get("why_now"):
                doc.add_paragraph(f"为何此时: {o['why_now']}")
            if o.get("entry_path"):
                doc.add_paragraph(f"切入路径: {o['entry_path']}")

    risks = brief.get("risks_and_threats") or []
    if risks:
        doc.add_heading("风险与威胁", level=2)
        for r in risks:
            extra = "，".join(
                x
                for x in [
                    f"可能性 {r['likelihood']}" if r.get("likelihood") else "",
                    f"缓解: {r['mitigation']}" if r.get("mitigation") else "",
                ]
                if x
            )
            doc.add_paragraph(
                r.get("risk", "") + (f"（{extra}）" if extra else ""),
                style="List Bullet",
            )

    rp = brief.get("recommended_positioning") or {}
    if rp:
        doc.add_heading("推荐定位", level=2)
        if rp.get("statement"):
            p = doc.add_paragraph()
            run = p.add_run(rp["statement"])
            run.bold = True
            run.font.size = Pt(12)
        if rp.get("differentiation_logic"):
            doc.add_paragraph(f"差异化逻辑: {rp['differentiation_logic']}")
        if rp.get("why_this_brand_can_own_it"):
            doc.add_paragraph(f"为何本品牌能占据: {rp['why_this_brand_can_own_it']}")
        for pp in rp.get("proof_points") or []:
            if isinstance(pp, dict) and pp.get("claim"):
                doc.add_paragraph(pp["claim"], style="List Bullet")


def _add_evidence_list(doc: Document, evidence: list[dict]) -> None:
    """添加论据列表"""
    if not evidence:
        return
    doc.add_paragraph("")
    doc.add_paragraph("数据论据:", style="Normal")
    for e in evidence:
        desc = e.get("description", "")
        etype = e.get("type", "")
        source = e.get("source", "")
        text = f"[{etype}] {desc}"
        if source:
            text += f" (来源: {source})"
        doc.add_paragraph(text, style="List Bullet")
